from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi import UploadFile
from pypdf import PdfReader

from app.core.config import settings


SUPPORTED_FILE_TYPES = {
    ".txt": "纯文本",
    ".md": "Markdown",
    ".pdf": "PDF",
    ".docx": "Word",
}


class FileExtractError(Exception):
    pass


def extract_text_from_upload(file: UploadFile) -> str:
    filename = file.filename or "未命名文件"
    suffix = Path(filename).suffix.lower()

    if suffix not in SUPPORTED_FILE_TYPES:
        supported = "、".join(SUPPORTED_FILE_TYPES)
        raise FileExtractError(f"暂不支持该文件类型，请上传 {supported} 文件。")

    content = file.file.read()
    max_upload_size = settings.max_upload_size_mb * 1024 * 1024
    if len(content) > max_upload_size:
        raise FileExtractError(
            f"文件不能超过 {settings.max_upload_size_mb} MB，请压缩后重试。"
        )

    if not content:
        raise FileExtractError("上传文件为空，请重新选择。")

    if suffix in {".txt", ".md"}:
        text = _decode_text(content)
    elif suffix == ".docx":
        text = _extract_docx_text(content)
    else:
        text = _extract_pdf_text(content)

    normalized = _normalize_text(text)
    if not normalized:
        raise FileExtractError("未从文件中解析到可用文本，请检查文件内容。")

    return normalized


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue

    raise FileExtractError("文本文件编码无法识别，请转成 UTF-8 后重试。")


def _extract_docx_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    chunks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))

    return "\n".join(chunks)


def _extract_pdf_text(content: bytes) -> str:
    reader = PdfReader(BytesIO(content))
    pages = []
    for page in reader.pages:
        extracted = page.extract_text() or ""
        if extracted.strip():
            pages.append(extracted)

    return "\n".join(pages)


def _normalize_text(value: str) -> str:
    lines = [line.rstrip() for line in value.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    trimmed = [line for line in lines]
    return "\n".join(trimmed).strip()
